"""
generate_report.py

This script will generate our report. 

It first reads the CSV of field inspection records, mirroring data coming from Survey123, then uses OpenAI to clean up
field notes into full sentences, then generates a small location map for
each site, and finally it produces a single Word doc file with one section per each
inspection. It is modeled on the Survey123 like data (in real project, API end point from Survey123/AGOL would supply the required data).

"""

import os
import csv
from pathlib import Path

from dotenv import load_dotenv
from openai import OpenAI
import folium
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- Config ----------
DATA_DIR = Path("data")
PHOTOS_DIR = DATA_DIR / "photos"
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)
MAPS_DIR = OUTPUT_DIR / "maps"
MAPS_DIR.mkdir(exist_ok=True)

CSV_PATH = DATA_DIR / "inspections.csv"
REPORT_PATH = OUTPUT_DIR / "inspection_report.docx"

OPENAI_MODEL = "gpt-4o-mini"  # Using gpt-4o-mini because it is quite cheap.

# ...................Setting up 
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# ....................Adding helpers
def clean_field_note(raw_note: str) -> str:
    """
    This function would send a raw field note to OpenAI and get back a single clean sentence.
    It would preserves numbers, units, and place names exactly as written, rest it will improve.
    """
    prompt = (
        "You clean up short field notes from infrastructure inspectors. "
        "Rewrite the note as one clear, complete English sentence. "
        "Keep all numbers, units, abbreviations (DI, TP, etc.), and place names "
        "exactly as written. Do not add information that is not in the note. "
        "Return only the cleaned sentence, no preamble."
    )
    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": raw_note},
        ],
        temperature=0.2,
    )
    return response.choices[0].message.content.strip()


def make_site_map(lat: float, lon: float, site_id: int) -> Path:
    """
    This function would generate a small PNG map file that shows the site location.
    It uses matplotlib for a simple offline map (a marker on a basic plot). I could've used Folium as well instead.
    """
    fig, ax = plt.subplots(figsize=(4, 3))
    ax.scatter([lon], [lat], c="red", s=200, marker="*", zorder=3)
    ax.set_xlim(lon - 0.02, lon + 0.02)
    ax.set_ylim(lat - 0.015, lat + 0.015)
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    ax.set_title(f"Site {site_id} location")
    ax.grid(True, alpha=0.3)
    ax.set_aspect("equal")

    map_path = MAPS_DIR / f"map_site_{site_id}.png"
    fig.savefig(map_path, dpi=120, bbox_inches="tight")
    plt.close(fig)
    return map_path


def add_inspection_section(doc: Document, record: dict, cleaned: str,
                           photo_path: Path, map_path: Path):
    """
    This function would append one formatted inspection section to the Word doc file.
    """
    # Section heading
    heading = doc.add_heading(
        f"Site {record['id']} — {record['creation_date']}", level=2
    )

    # Metadata table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "Inspector"
    table.cell(0, 1).text = record["creator"]
    table.cell(1, 0).text = "Coordinates"
    table.cell(1, 1).text = f"{record['lat']}, {record['lon']}"
    table.cell(2, 0).text = "Creation date"
    table.cell(2, 1).text = record["creation_date"]

    # Cleaned narrative
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Inspector observation:")
    run.bold = True
    doc.add_paragraph(cleaned)

    # Photo with map side by side via a 1 row, 2 column table
    doc.add_paragraph()
    media_table = doc.add_table(rows=2, cols=2)
    media_table.cell(0, 0).text = "Site photo"
    media_table.cell(0, 1).text = "Location map"
    photo_cell = media_table.cell(1, 0).paragraphs[0]
    photo_run = photo_cell.add_run()
    photo_run.add_picture(str(photo_path), width=Inches(2.8))
    map_cell = media_table.cell(1, 1).paragraphs[0]
    map_run = map_cell.add_run()
    map_run.add_picture(str(map_path), width=Inches(2.8))

    # Adding spacing between differ sections
    doc.add_paragraph()
    doc.add_paragraph()


# ...................Main pipeline 
def main():
    print(f"Reading inspections from {CSV_PATH}")
    with open(CSV_PATH, encoding="utf-8") as f:
        records = list(csv.DictReader(f))
    print(f"Found {len(records)} inspections")

    # Casting lat/lon to floats for better working
    for r in records:
        r["lat"] = float(r["lat"])
        r["lon"] = float(r["lon"])

    # Now building the Word doc file
    doc = Document()
    title = doc.add_heading("Field Inspection Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "This report is auto generated from Survey123 mirroring field submissions. "
        "The field notes are cleaned using OpenAI gpt-4o-mini, whilst maps are generated locally."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Processing each inspection
    for i, record in enumerate(records, start=1):
        print(f"\n--- Processing item {i} of {len(records)} (Site {record['id']}) ---")

        # Step A: Cleaning the note
        print(f"  Original: {record['notes']}")
        cleaned = clean_field_note(record["notes"])
        print(f"  Cleaned:  {cleaned}")

        # Step B: Making the map 
        map_path = make_site_map(record["lat"], record["lon"], record["id"])
        print(f"  Map saved to {map_path}")

        # Step C: Confirm photo exists, just to make sure
        photo_path = PHOTOS_DIR / record["photo"]
        if not photo_path.exists():
            print(f"  WARNING: photo {photo_path} missing, skipping image")
            continue

        # Step D: Adding to doc now
        add_inspection_section(doc, record, cleaned, photo_path, map_path)

    # Save the report
    doc.save(REPORT_PATH)
    print(f"\nReport saved: {REPORT_PATH.resolve()}")


if __name__ == "__main__":
    main()